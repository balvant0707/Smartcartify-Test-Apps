from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "docs/CartLift-Policies-and-FAQs.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.add_run(text)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.add_run(text)
    return p


def add_faq(doc, question, answer):
    p = doc.add_paragraph(style="Heading 2")
    p.add_run(question)
    add_body(doc, answer)


def apply_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ["List Bullet", "List Number"]:
        style = styles[list_style]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("CartLift: Cart Drawer & Upsell | Policies and FAQs")
    set_run_font(r, size=9, color="666666")


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("CartLift: Cart Drawer & Upsell")
    set_run_font(r, size=24, color="0B2545", bold=True)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(12)
    r2 = p2.add_run("App Policies and Frequently Asked Questions")
    set_run_font(r2, size=14, color="555555")
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    r3 = meta.add_run(
        "Prepared from the current app codebase and docs. Last updated: June 26, 2026."
    )
    set_run_font(r3, size=10, color="666666")


def add_info_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Field"
    table.cell(0, 1).text = "Details"
    rows = [
        ("App name", "CartLift: Cart Drawer & Upsell"),
        ("Developer", "Pryxo Tech Private Limited"),
        ("Primary privacy contact", "balvant@pryxotech.com"),
        ("Merchant support email", "info@pryxotech.com; support@pryxotech.com"),
        ("Support links", "https://cartmilestone.smartreminder.in/ and https://pryxotech.com/#inquiry-now"),
    ]
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    set_table_widths(table, [2700, 6660])
    set_repeat_table_header(table.rows[0])
    for row_idx, row in enumerate(table.rows):
        if row_idx == 0:
            for cell in row.cells:
                set_cell_shading(cell, "E8EEF5")
        else:
            set_cell_shading(row.cells[0], "F2F4F7")
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
            for run in cell.paragraphs[0].runs:
                set_run_font(run, size=10)
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True


def add_plan_table(doc):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Plan", "Price", "Trial", "Included rule access"]
    for i, text in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = text
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    rows = [
        ("Free", "$0", "No trial", "One shipping rule, one automation discount rule, customize settings, cart step progress and preview."),
        ("Professional Monthly", "$5 every 30 days", "7 days", "Unlimited shipping rules, discount rules, free product rules, Buy X Get Y rules, customization, cart step progress and preview."),
        ("Professional Yearly", "$49 annually", "7 days", "Same Professional feature access as monthly, billed annually."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
    set_table_widths(table, [1700, 1900, 1300, 4460])
    set_repeat_table_header(table.rows[0])
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, size=9.5)


def build_doc():
    doc = Document()
    apply_styles(doc)
    add_footer(doc)
    add_title(doc)
    add_info_table(doc)

    add_heading(doc, "Purpose", 1)
    add_body(
        doc,
        "This document summarizes the merchant-facing policies, data practices, support channels, billing behavior, and frequently asked questions for CartLift: Cart Drawer & Upsell.",
    )
    add_body(
        doc,
        "It is intended as a working policy and FAQ reference for app listing, support, onboarding, and internal review. Legal counsel should review it before public use as a binding legal policy.",
    )

    add_heading(doc, "App Overview", 1)
    add_body(
        doc,
        "CartLift is a Shopify cart drawer and promotion app that helps merchants increase conversion and average order value through cart progress messaging, discounts, free gifts, Buy X Get Y offers, upsell recommendations, and cart drawer customization.",
    )
    add_bullet(doc, "Cart Goal: shows tiered cart goals and progress bars based on cart value or quantity.")
    add_bullet(doc, "Shipping Rule: offers free or discounted shipping based on cart thresholds.")
    add_bullet(doc, "Automatic Discount: applies fixed or percentage discounts when spend or quantity rules are met.")
    add_bullet(doc, "Free Product Discount: adds or presents a free gift when customers meet configured thresholds.")
    add_bullet(doc, "Code Discount: displays discount code campaigns inside the cart drawer or announcement area.")
    add_bullet(doc, "Buy X Get Y Discount: rewards qualifying purchases using Shopify discount behavior.")
    add_bullet(doc, "Upsell Product Rules: recommends related or manually selected products in the cart drawer.")
    add_bullet(doc, "Customize & Preview: lets merchants adjust cart drawer colors, fonts, icons, layout, and preview settings.")

    add_heading(doc, "Merchant Data Policy", 1)
    add_heading(doc, "Information Collected", 2)
    add_body(doc, "When a merchant installs and uses the app, the app may collect and store shop-level and merchant-level information needed to operate the Shopify app.")
    add_bullet(doc, "Shop domain, primary domain, installation state, onboarding details, and timestamps.")
    add_bullet(doc, "Store owner or administrator name, email address, phone number, country, city, currency, and related contact fields when provided by Shopify.")
    add_bullet(doc, "Shopify session metadata, scopes, locale, expiration, and access token needed to operate the app.")
    add_bullet(doc, "Merchant-created app configuration, including rules, style settings, cart drawer preferences, selected product or collection identifiers, and upsell settings.")
    add_bullet(doc, "Billing and subscription data received from Shopify, including plan ID, plan name, status, billing interval, amount, currency, trial days, subscription ID, and current period end.")

    add_heading(doc, "Customer and Storefront Data", 2)
    add_body(
        doc,
        "The storefront script reads limited cart information in real time to display cart drawer content, progress messaging, offers, gifts, discounts, and upsell recommendations. Cart data is processed dynamically and is not stored as customer profiles by the app.",
    )
    add_body(doc, "The app does not store customer personal data such as customer names, email addresses, postal addresses, or payment information.")

    add_heading(doc, "How Data Is Used", 2)
    add_bullet(doc, "Authenticate the merchant and operate the app inside Shopify.")
    add_bullet(doc, "Save and render merchant-configured rules, settings, styles, and storefront customizations.")
    add_bullet(doc, "Retrieve relevant Shopify product, collection, discount, theme, and billing data required by app features.")
    add_bullet(doc, "Manage plan status, billing redirects, trial information, cancellation status, and subscription webhooks.")
    add_bullet(doc, "Send installation, uninstallation, onboarding, or support-related emails.")
    add_bullet(doc, "Maintain security, troubleshoot issues, and comply with Shopify platform requirements.")

    add_heading(doc, "Data Sharing Policy", 2)
    add_body(doc, "The app does not sell, rent, or trade merchant data or customer data.")
    add_body(
        doc,
        "Limited information may be shared with service providers only when required to operate the app, including Shopify for platform APIs, billing, authentication, and webhooks; infrastructure and database providers; and email delivery providers.",
    )

    add_heading(doc, "Browser Storage and Tracking", 2)
    add_body(
        doc,
        "The app does not use its own storefront advertising trackers. The storefront script may use browser sessionStorage for limited temporary behavior, such as remembering a pending discount code during the browsing session. That browser-session value is not stored on the app server by that feature.",
    )

    add_heading(doc, "Retention and Deletion Policy", 1)
    add_body(doc, "Merchant configuration data is retained while the app is installed and active, subject to operational and legal needs.")
    add_body(doc, "When Shopify sends a shop/redact privacy webhook, the app deletes shop-related app data, including configured shipping rules, discount rules, free gift rules, Buy X Get Y rules, style settings, upsell settings, plan subscription records, Shopify sessions, and the shop record.")
    add_body(doc, "For Shopify customer data access and customer redaction webhooks, the app responds that it does not store customer personal data in its own database.")

    add_heading(doc, "Security Policy", 1)
    add_bullet(doc, "Shopify OAuth and admin authentication are used for merchant access.")
    add_bullet(doc, "Shopify access tokens are encrypted at rest according to the app privacy policy.")
    add_bullet(doc, "HTTPS/TLS is used for communication between the app, Shopify, and merchant browsers.")
    add_bullet(doc, "Shopify webhooks and app proxy requests are verified before processing.")
    add_bullet(doc, "Merchant data is used only for app operation, support, billing, security, and compliance purposes.")

    add_heading(doc, "Billing and Plan Policy", 1)
    add_body(doc, "The app supports a Free plan and two Professional paid plans. Paid billing is handled through Shopify Billing.")
    add_plan_table(doc)
    add_body(doc, "Paid plans are requested through Shopify billing and are initially recorded as pending until Shopify confirms activation. Merchants can cancel the current paid plan from the pricing page. Paid subscription cancellation is sent through Shopify billing with prorate set to false in the current implementation.")

    add_heading(doc, "Support Policy", 1)
    add_body(doc, "Merchants should include their shop URL, steps to reproduce, and screenshots when contacting support. The app help page links to Pryxo Tech inquiry support and provides an email contact.")
    add_bullet(doc, "Support email shown in app: info@pryxotech.com.")
    add_bullet(doc, "Support email used by app emails when no environment override is set: support@pryxotech.com.")
    add_bullet(doc, "Privacy contact in the existing privacy policy: balvant@pryxotech.com.")
    add_bullet(doc, "Support URL in the existing privacy policy: https://cartmilestone.smartreminder.in/.")
    add_bullet(doc, "Inquiry support link in the app: https://pryxotech.com/#inquiry-now.")

    add_heading(doc, "Frequently Asked Questions", 1)
    add_faq(doc, "What does CartLift do?", "CartLift replaces or supplements the storefront cart experience with a customizable cart drawer and promotional tools such as cart goals, automatic discounts, discount codes, Buy X Get Y offers, free gifts, shipping progress, and upsell products.")
    add_faq(doc, "How do I enable the cart drawer on my store?", "Open the app dashboard and use Open App Embeds to enable the CartLift app embed in the Shopify theme editor. The dashboard shows whether the app embed is ON or OFF.")
    add_faq(doc, "Does the app store customer personal data?", "No. The app is designed not to store customer names, email addresses, postal addresses, or payment information. Storefront cart data is processed in real time for display and promotion logic.")
    add_faq(doc, "What data is deleted when a shop is redacted?", "The app deletes shop-related configuration and records, including shipping rules, discount rules, free gift rules, Buy X Get Y rules, style settings, upsell settings, plan subscription records, sessions, and the shop record.")
    add_faq(doc, "Does the app use Shopify billing?", "Yes. Professional Monthly and Professional Yearly are requested through Shopify Billing. The app records subscription status and related billing fields so the pricing page can show the current plan state.")
    add_faq(doc, "Is there a free trial?", "Yes. The Professional Monthly and Professional Yearly plans are configured with a 7 day free trial. The Free plan has no trial because it costs $0.")
    add_faq(doc, "Can merchants cancel a paid plan?", "Yes. The pricing page includes a Cancel current plan action when a paid subscription is active. The app sends cancellation to Shopify Billing and updates the local subscription record.")
    add_faq(doc, "What features are limited on the Free plan?", "The Free plan includes priority support, one shipping rule, one automation discount rule, customize settings, and cart step progress and preview. Professional plans unlock unlimited shipping, discount, free product, and Buy X Get Y rules.")
    add_faq(doc, "Can the cart drawer be customized?", "Yes. Merchants can customize colors, fonts, button styles, backgrounds, cart icon type, checkout button text, drawer position, mobile layout, sticky checkout, and related display settings.")
    add_faq(doc, "Can merchants upload custom cart icons?", "Yes. The customize page supports cart icon uploads for PNG, JPG, WebP, GIF, and SVG files up to 2 MB.")
    add_faq(doc, "Which campaign types are available?", "The app includes Cart Goal, Shipping Rule, Automatic Discount, Free Product Discount, Code Discount, Buy X Get Y Discount, Upsell Product Rules, and Customize & Preview.")
    add_faq(doc, "How do upsell rules work?", "Upsell rules display recommended products inside the cart drawer. Merchants can use automatic Shopify-related recommendations or manually select products and collections.")
    add_faq(doc, "How should a merchant contact support?", "Use the Help & Support page in the app, send email to info@pryxotech.com, or use the inquiry support link at https://pryxotech.com/#inquiry-now. Include shop URL, steps, and screenshots.")
    add_faq(doc, "Does the app send emails?", "Yes. The app includes installation and uninstallation email templates for merchants and internal notifications. Support-related replies direct merchants to the configured support email.")
    add_faq(doc, "Is this document a legal agreement?", "No. It is a repo-derived working policy and FAQ reference. Review with legal counsel before publishing it as binding terms or a public legal policy.")

    add_heading(doc, "Source Notes", 1)
    add_body(doc, "This document was prepared from docs/PRIVACY_POLICY.md, app/routes/webhooks.gdpr.jsx, app/lib/plans.js, app/routes/app.pricing.jsx, app/routes/app.help.jsx, app/routes/app.campaigns.jsx, app/routes/app.customize-preview.jsx, and app/lib/emailTemplates.server.js.")

    doc.core_properties.title = "CartLift Policies and FAQs"
    doc.core_properties.subject = "Policies, data handling, billing, support, and FAQ reference"
    doc.core_properties.author = "Pryxo Tech Private Limited"
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
